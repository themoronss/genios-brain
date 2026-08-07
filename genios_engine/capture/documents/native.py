from __future__ import annotations

import io
from html.parser import HTMLParser

from .base import OCR_MIN_CONFIDENCE, DocumentInput, DocumentResult, OcrEngine
from .router import route_document

# Native text extraction — NO OCR. If a document already has a text layer (HTML, digital
# PDF, docx, txt/md), we pull it straight out; Tesseract is only the fallback for scanned
# images with no text layer.

_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class _HTMLText(HTMLParser):
    _SKIP = {"script", "style", "head"}

    def __init__(self) -> None:
        super().__init__()
        self._parts: list[str] = []
        self._skip = 0

    def handle_starttag(self, tag, attrs):
        if tag in self._SKIP:
            self._skip += 1

    def handle_endtag(self, tag):
        if tag in self._SKIP and self._skip:
            self._skip -= 1

    def handle_data(self, data):
        if not self._skip and data.strip():
            self._parts.append(data.strip())

    def text(self) -> str:
        return " ".join(self._parts)


def _html_to_text(s: str) -> str:
    p = _HTMLText()
    p.feed(s)
    return p.text()


def _docx_to_text(raw: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(raw))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for tbl in doc.tables:
        for row in tbl.rows:
            parts += [c.text.strip() for c in row.cells if c.text.strip()]
    return "\n".join(parts)


def _pdf_to_text(raw: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(raw))
    return "\n".join((pg.extract_text() or "") for pg in reader.pages).strip()


def extract_native_text(*, mime: str, data: bytes | str, filename: str = "") -> str | None:
    """Return the document's text layer, or None if it has none (scanned image →
    OCR fallback) or the format is unsupported. Never raises."""
    mime = (mime or "").lower()
    name = (filename or "").lower()
    raw = data.encode() if isinstance(data, str) else data
    txt = data if isinstance(data, str) else data.decode(errors="ignore")
    try:
        if mime in ("text/plain", "text/markdown") or name.endswith((".txt", ".md")):
            return txt
        if mime == "text/html" or name.endswith((".html", ".htm")):
            return _html_to_text(txt)
        if mime == _DOCX or name.endswith(".docx"):
            return _docx_to_text(raw)
        if mime == "application/pdf" or name.endswith(".pdf"):
            return _pdf_to_text(raw)
    except Exception:
        return None
    return None


def _ocr_image_bytes(mime: str, data: bytes, filename: str, ocr: OcrEngine) -> DocumentResult:
    """Materialise raw image bytes to a short-lived temp file and OCR them — route_document's OCR
    branch needs an image_ref, and before this no caller ever set one, so OCR was dormant."""
    import os
    import tempfile
    suffix = "." + ((mime or "").split("/", 1)[-1] or "img").split(";")[0]
    fd, path = tempfile.mkstemp(suffix=suffix)
    try:
        with os.fdopen(fd, "wb") as f:
            f.write(data)
        return route_document(DocumentInput(mime=mime, filename=filename, image_ref=path), ocr=ocr)
    finally:
        try:
            os.remove(path)
        except OSError:
            pass


def _ocr_pdf_bytes(data: bytes, ocr: OcrEngine) -> DocumentResult | None:
    """Rasterize a scanned PDF's pages (poppler via pdf2image) and OCR each, concatenating the text.
    Returns None if pdf2image/poppler is unavailable OR nothing was read, so the caller falls back
    to 'unsupported' gracefully — a missing OCR toolchain must never crash ingestion."""
    try:
        from pdf2image import convert_from_bytes
    except Exception:                                    # pdf2image/poppler not installed
        return None
    try:
        images = convert_from_bytes(data)
    except Exception:                                    # corrupt PDF / poppler failure
        return None
    import os
    import tempfile
    texts: list[str] = []
    confs: list[float] = []
    for img in images:
        fd, path = tempfile.mkstemp(suffix=".png")
        try:
            with os.fdopen(fd, "wb") as f:
                img.save(f, "PNG")
            r = ocr.ocr(path)
            if r.text:
                texts.append(r.text)
            confs.append(r.avg_confidence)
        finally:
            try:
                os.remove(path)
            except OSError:
                pass
    if not texts:
        return None
    avg = sum(confs) / len(confs) if confs else 0.0
    status = "accepted" if avg >= OCR_MIN_CONFIDENCE else "ocr_review_required"
    return DocumentResult(text="\n".join(texts), native_parse_used=False, ocr_used=True,
                          ocr_engine=getattr(ocr, "name", None), ocr_pages=len(texts),
                          avg_confidence=avg, status=status)


def process_document(*, mime: str, data: bytes | str, filename: str = "",
                     image_ref: str | None = None,
                     ocr: OcrEngine | None = None) -> DocumentResult:
    """Full path: native text if the format has it, else OCR (if wired), else unsupported.

    OCR now actually RUNS when an engine is wired (before, no caller set `image_ref`, so it was
    dormant engine-wide): an image's bytes are materialised to a temp file, and a scanned PDF's
    pages are rasterized (poppler) and OCR'd. A missing OCR toolchain falls back to 'unsupported',
    never a crash."""
    text = extract_native_text(mime=mime, data=data, filename=filename)
    if (image_ref is None and ocr is not None and not (text and text.strip())
            and isinstance(data, (bytes, bytearray))):
        lower = (mime or "").lower()
        if lower.startswith("image/"):
            return _ocr_image_bytes(mime, data, filename, ocr)
        if lower == "application/pdf" or filename.lower().endswith(".pdf"):
            pdf_result = _ocr_pdf_bytes(data, ocr)
            if pdf_result is not None:
                return pdf_result                        # else: poppler absent → fall through to unsupported
    doc = DocumentInput(mime=mime, filename=filename, text_layer=text, image_ref=image_ref)
    return route_document(doc, ocr=ocr)


# text-ish formats we decode straight to utf-8 when there is no modelled parser (csv/json/logs/yaml).
# NOT html — that must be tag-stripped (extract_native_text), never returned raw.
_PLAINTEXT_EXTS = ("txt", "md", "markdown", "csv", "tsv", "json", "log", "yaml", "yml")


def extract_text_best_effort(*, mime: str, data: bytes | str, filename: str = "",
                             ocr: OcrEngine | None = None) -> str:
    """Best-effort document text for the intake doors (dashboard uploads): the full native (+OCR
    when wired) path — IDENTICAL to email/Drive attachments — then a stripped-native recovery for
    short text layers, then a utf-8 fallback for plain-text formats with no modelled parser. Returns
    "" (never None) for a binary/scanned file that yields nothing, so the caller can honestly report
    "no extractable text" instead of indexing decoded garbage. (The upload door was pypdf-only before
    this — inconsistent with email/Drive, and it utf-8-decoded binaries into garbage.)"""
    r = process_document(mime=mime or "", data=data, filename=filename, ocr=ocr)
    if r.text and r.text.strip():
        return r.text
    # process_document rejects a native text layer under _MIN_NATIVE_CHARS and formats it doesn't
    # model — recover a stripped layer (short html, tiny pdf/docx) before the raw-decode fallback.
    native = extract_native_text(mime=mime or "", data=data, filename=filename)
    if native and native.strip():
        return native
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext in _PLAINTEXT_EXTS:
        return data.decode("utf-8", errors="ignore") if isinstance(data, (bytes, bytearray)) else str(data)
    return ""
