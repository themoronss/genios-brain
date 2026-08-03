"""Generate GeniOS MCP Integration Guide docx."""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Styles ───────────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x16, 0x21, 0x3E)
    return p

def h3(text):
    return doc.add_heading(text, level=3)

def para(text, bold=False):
    p = doc.add_paragraph(text)
    if bold:
        for run in p.runs:
            run.bold = True
    return p

def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")

def code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x2E)
    p.paragraph_format.left_indent = Inches(0.3)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F3F4F6")
    p._element.get_or_add_pPr().append(shading)
    return p

def divider():
    doc.add_paragraph("─" * 72)

# ── Cover ────────────────────────────────────────────────────────────────────
title = doc.add_heading("GeniOS MCP Integration Guide", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x0F, 0x0F, 0x23)

sub = doc.add_paragraph("Connect any AI client to your relationship brain")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x77)

doc.add_paragraph()
divider()
doc.add_paragraph()

# ── Overview ─────────────────────────────────────────────────────────────────
h1("Overview")
para(
    "GeniOS exposes your relationship memory as an MCP (Model Context Protocol) server. "
    "Once connected, any AI client — Claude, ChatGPT, Cursor, VS Code — can pull "
    "live contact context, relationship alerts, and open commitments directly into every conversation."
)
doc.add_paragraph()

para("There are two connection methods:", bold=True)

table = doc.add_table(rows=3, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = ""
hdr[1].text = "Remote MCP (HTTPS)"
hdr[2].text = "Local MCP (stdio)"
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True

r1 = table.rows[1].cells
r1[0].text = "Best for"
r1[1].text = "Claude.ai, ChatGPT, mobile"
r1[2].text = "Claude Desktop, Cursor, VS Code, Cline, Claude Code"

r2 = table.rows[2].cells
r2[0].text = "What you install"
r2[1].text = "Nothing — just a URL + API key"
r2[2].text = "Node.js 18+ and the @genios/mcp package"

doc.add_paragraph()

# ── Step 1: API Key ──────────────────────────────────────────────────────────
h1("Step 1 — Get your API key")
bullet("Open GeniOS → Settings → API Keys")
bullet("Click Generate key")
bullet("Copy the key — it starts with gn_live_…")
bullet("Treat it like a password. Never commit it to version control.")
doc.add_paragraph()

# ── Step 2: Remote ───────────────────────────────────────────────────────────
h1("Step 2 — Remote MCP (Claude.ai / ChatGPT / online)")
para(
    "Remote MCP connects directly to the GeniOS backend over HTTPS. "
    "No installation needed — just a URL and your API key."
)
doc.add_paragraph()

para("Server URL:", bold=True)
code("https://squid-app-2zuqf.ondigitalocean.app/mcp")
doc.add_paragraph()

h2("Claude.ai")
bullet("Open claude.ai → Settings → Connectors → Add custom connector")
bullet("Name: GeniOS")
bullet("Server URL: paste the /mcp URL above")
bullet("Advanced settings → Authorization → Bearer token → paste your gn_live_… key")
bullet("Click Add, then Connect")
bullet("Open a new chat — look for the 🔧 icon to confirm tools loaded")
doc.add_paragraph()

h2("ChatGPT (Pro / Business / Enterprise)")
bullet("Settings → Connectors → Add custom connector")
bullet("Paste the same URL and Bearer token")
bullet("Save and refresh the chat")
doc.add_paragraph()

h2("Verify the connection")
para("Run this in a terminal to confirm the server is responding:")
code(
    'curl -X POST https://squid-app-2zuqf.ondigitalocean.app/mcp \\\n'
    '  -H "Authorization: Bearer gn_live_YOUR_KEY" \\\n'
    '  -H "Content-Type: application/json" \\\n'
    '  -d \'{"jsonrpc":"2.0","id":1,"method":"tools/list"}\''
)
para('Expected: a JSON list of 10 genios_* tools. If you get 401, your key is wrong.')
doc.add_paragraph()

# ── Step 3: Local ────────────────────────────────────────────────────────────
h1("Step 3 — Local MCP (Claude Desktop / Cursor / Claude Code)")

h2("Dependencies")
para("You need Node.js 18 or newer. Check:")
code("node --version")
para("If not installed:")
bullet("Windows/macOS: download from nodejs.org")
bullet("Linux (Ubuntu/Debian): sudo apt install nodejs")
doc.add_paragraph()

h2("Install the package")
code("npm install -g @genios/mcp")
doc.add_paragraph()

h2("Set environment variables")
para("Create a .env file or set these in your shell:")
code(
    "GENIOS_API_KEY=gn_live_your_key_here\n"
    "GENIOS_BASE_URL=https://squid-app-2zuqf.ondigitalocean.app\n"
    "GENIOS_AGENT_ID=yourname-mcp"
)
doc.add_paragraph()

h2("Wire into Claude Desktop")
para("Edit the config file for your OS:")
bullet("macOS: ~/Library/Application Support/Claude/claude_desktop_config.json")
bullet("Linux: ~/.config/Claude/claude_desktop_config.json")
bullet("Windows: %APPDATA%\\Claude\\claude_desktop_config.json")
doc.add_paragraph()
para("Add this block (keep any existing servers):")
code(
    '{\n'
    '  "mcpServers": {\n'
    '    "genios": {\n'
    '      "command": "npx",\n'
    '      "args": ["-y", "@genios/mcp"],\n'
    '      "env": {\n'
    '        "GENIOS_API_KEY": "gn_live_your_key_here",\n'
    '        "GENIOS_BASE_URL": "https://squid-app-2zuqf.ondigitalocean.app"\n'
    '      }\n'
    '    }\n'
    '  }\n'
    '}'
)
para("Fully quit Claude Desktop (File → Quit), then reopen. Tools appear in the 🔧 icon.")
doc.add_paragraph()

h2("Wire into Claude Code")
code(
    'claude mcp add genios \\\n'
    '  -e GENIOS_API_KEY=gn_live_your_key_here \\\n'
    '  -e GENIOS_BASE_URL=https://squid-app-2zuqf.ondigitalocean.app \\\n'
    '  -- npx -y @genios/mcp'
)
doc.add_paragraph()

h2("Wire into Cursor / VS Code / Cline")
para("Same JSON block as Claude Desktop — paste it in the MCP settings panel of each tool.")
doc.add_paragraph()

# ── Tools ────────────────────────────────────────────────────────────────────
h1("Tools available")

table2 = doc.add_table(rows=11, cols=3)
table2.style = "Table Grid"
hdr2 = table2.rows[0].cells
hdr2[0].text = "Tool"
hdr2[1].text = "What it does"
hdr2[2].text = "Writes data?"
for cell in hdr2:
    cell.paragraphs[0].runs[0].bold = True

tools = [
    ("genios_search_contacts",   "Search or list contacts by name, email, stage, or silence window", "No"),
    ("genios_get_context",       "Full relationship memory — stage, sentiment, commitments, guidance", "No"),
    ("genios_list_segments",     "List all contact segments",                                         "No"),
    ("genios_get_segment_members","Contacts inside a specific segment",                               "No"),
    ("genios_org_info",          "Plan tier, usage stats, graph totals",                              "No"),
    ("genios_list_insights",     "Proactive alerts — cooling contacts, overdue commitments, anomalies","No"),
    ("genios_sync_status",       "Freshness of Gmail / Calendar / Slack sync",                       "No"),
    ("genios_log_interaction",   "Record an outreach action on a contact",                           "Yes"),
    ("genios_log_outcome",       "Feedback on whether context was useful",                           "Yes"),
    ("genios_trigger_scan",      "Force a fresh proactive scan now",                                 "Yes"),
]
for i, (name, desc, writes) in enumerate(tools):
    row = table2.rows[i + 1].cells
    row[0].text = name
    row[1].text = desc
    row[2].text = writes

doc.add_paragraph()

# ── Test prompts ─────────────────────────────────────────────────────────────
h1("Quick test prompts")
para("Drop these into any MCP-connected chat to verify tools are working:")
bullet('"What are my relationship priorities today?" → genios_list_insights')
bullet('"Give me context on alice@acmecorp.com" → genios_get_context')
bullet('"What segments do I have?" → genios_list_segments')
bullet('"How fresh is my Gmail sync?" → genios_sync_status')
bullet('"What plan am I on?" → genios_org_info')
doc.add_paragraph()

# ── Troubleshooting ──────────────────────────────────────────────────────────
h1("Troubleshooting")

table3 = doc.add_table(rows=9, cols=2)
table3.style = "Table Grid"
hdr3 = table3.rows[0].cells
hdr3[0].text = "Symptom"
hdr3[1].text = "Fix"
for cell in hdr3:
    cell.paragraphs[0].runs[0].bold = True

issues = [
    ("Tools don't appear in Claude Desktop", "File → Quit (don't just close window), reopen"),
    ("401 Unauthorized", "Key must be Authorization: Bearer gn_live_… — check spelling"),
    ("404 on /mcp", "Backend not deployed with MCP route — redeploy genios-brain"),
    ("Claude shows fake data / hallucinations", "Update @genios/mcp to latest: npm install -g @genios/mcp"),
    ("npx command not found", "Node.js not installed — see Step 3 dependencies"),
    ("Claude.ai 'Connection failed'", "URL must be https://, not http://"),
    ("Tool hangs then errors", "GENIOS_BASE_URL unreachable — check the URL"),
    ("ChatGPT doesn't see tool changes", "Remove connector and re-add — ChatGPT caches tool list"),
]
for i, (symptom, fix) in enumerate(issues):
    row = table3.rows[i + 1].cells
    row[0].text = symptom
    row[1].text = fix

doc.add_paragraph()

# ── Security ─────────────────────────────────────────────────────────────────
h1("Security")
bullet("Your gn_live_… key is the sole credential — treat it like a password")
bullet("Rotate: generate new key in GeniOS → update all clients → delete old key")
bullet("Every MCP call is rate-limited and audit-logged per org")
bullet("Never share your key or commit it to version control")

# ── Save ─────────────────────────────────────────────────────────────────────
out = Path(__file__).resolve().parent / "GeniOS_MCP_Integration_Guide.docx"
doc.save(str(out))
print(f"Saved: {out}")
